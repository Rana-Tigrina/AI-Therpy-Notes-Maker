"""
Document Generation Module for Clinical Therapy Notes.

Constructs formatted Microsoft Word (.docx) clinical reports containing
client credentials, therapist metadata, and structured clinical observations.
"""

from typing import Dict, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config.config import logger


class TherapyNotesDocumentGenerator:
    """
    Generates structured Word documents (.docx) for clinical therapy notes.
    """

    def __init__(self):
        self.document = Document()
        self._setup_document_margins()

    def _setup_document_margins(self):
        """Configure 1-inch standard clinical documentation margins."""
        for section in self.document.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_first_page(self):
        """Build initial credentials and session metadata page."""
        # Client Credentials Section
        heading = self.document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run("CLIENT CREDENTIALS")
        heading_run.bold = True
        heading_run.font.size = Pt(12)
        self.document.add_paragraph()

        client_fields = [
            ("NAME:", "AGE:"),
            ("GENDER:", "DATE OF BIRTH:"),
            ("RESIDENCE:", "CONTACT:"),
            ("E-MAIL:", "EMERGENCY CONTACT:"),
            ("CLIENT ID:", "APPOINTMENT ID:"),
        ]
        for left, right in client_fields:
            p = self.document.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(3))
            p.add_run(left).bold = True
            p.add_run("\t")
            p.add_run(right).bold = True

        self.document.add_paragraph()

        # Therapist Credentials Section
        therapist_heading = self.document.add_paragraph()
        therapist_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        therapist_run = therapist_heading.add_run("THERAPIST CREDENTIALS")
        therapist_run.bold = True
        therapist_run.font.size = Pt(12)
        self.document.add_paragraph()

        therapist_fields = [
            ("NAME:", "THERAPIST ID:"),
            ("RCI LICENSE NO:", "DESIGNATION:"),
            ("MENTOR:", "MENTOR LICENSE NO:"),
            ("Session No.:", ""),
            ("Session Date:", ""),
        ]
        for left, right in therapist_fields:
            p = self.document.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(3))
            p.add_run(left).bold = True
            if right:
                p.add_run("\t")
                p.add_run(right).bold = True

        self.document.add_page_break()

    def _add_therapy_notes_page(self, therapy_notes: Dict[str, Any]):
        """Render clinical notes sections into the document."""
        heading = self.document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run("Therapy Notes & Clinical Documentation")
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        self.document.add_paragraph()

        for section_name, content in therapy_notes.items():
            formatted_title = section_name.replace("_", " ").title()
            h = self.document.add_heading(level=1)
            run = h.add_run(formatted_title)
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(16, 44, 87)

            if isinstance(content, dict):
                for key, value in content.items():
                    sub_title = key.replace("_", " ").title()
                    p = self.document.add_paragraph()
                    p.add_run(f"{sub_title}: ").bold = True

                    if isinstance(value, list):
                        for item in value:
                            sub_p = self.document.add_paragraph()
                            sub_p.style = "List Bullet"
                            sub_p.add_run(str(item))
                    elif isinstance(value, dict):
                        for sub_k, sub_v in value.items():
                            sub_p = self.document.add_paragraph()
                            sub_p.style = "List Bullet"
                            sub_p.add_run(f"{sub_k}: {sub_v}")
                    else:
                        p.add_run(str(value))
            elif isinstance(content, list):
                for item in content:
                    p = self.document.add_paragraph()
                    p.style = "List Bullet"
                    p.add_run(str(item))
            else:
                p = self.document.add_paragraph()
                p.add_run(str(content))

            self.document.add_paragraph()

    def generate_document(self, therapy_notes: Dict[str, Any]) -> Document:
        """
        Generate complete clinical document from notes.

        Args:
            therapy_notes (dict): Structured clinical notes.

        Returns:
            docx.Document: The populated Word document object.
        """
        try:
            self._add_first_page()
            self._add_therapy_notes_page(therapy_notes)
            return self.document
        except Exception as e:
            logger.error(f"Error generating Word document: {e}")
            return self.document

    def save_document(self, filename: str):
        """Save the document to specified file path."""
        try:
            self.document.save(filename)
        except Exception as e:
            logger.error(f"Error saving DOCX to {filename}: {e}")
            raise